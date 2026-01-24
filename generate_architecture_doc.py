"""Generate EconStats Architecture Word Document."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

def add_heading_with_color(doc, text, level, color=None):
    """Add a heading with optional color."""
    heading = doc.add_heading(text, level=level)
    if color:
        for run in heading.runs:
            run.font.color.rgb = color
    return heading

def add_step_box(doc, step_num, title, what_it_does, how_it_works, example, file_ref):
    """Add a formatted step explanation."""
    # Step header
    p = doc.add_paragraph()
    run = p.add_run(f"Step {step_num}: {title}")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(37, 99, 235)  # Blue

    # What it does
    p = doc.add_paragraph()
    run = p.add_run("What it does: ")
    run.bold = True
    p.add_run(what_it_does)

    # How it works
    p = doc.add_paragraph()
    run = p.add_run("How it works: ")
    run.bold = True
    p.add_run(how_it_works)

    # Example
    if example:
        p = doc.add_paragraph()
        run = p.add_run("Example: ")
        run.bold = True
        run.font.color.rgb = RGBColor(22, 163, 74)  # Green
        p.add_run(example)

    # File reference
    p = doc.add_paragraph()
    run = p.add_run("File: ")
    run.bold = True
    run.font.color.rgb = RGBColor(107, 114, 128)  # Gray
    code_run = p.add_run(file_ref)
    code_run.font.name = "Courier New"
    code_run.font.size = Pt(10)

    doc.add_paragraph()  # Spacing

def create_architecture_doc(filename):
    """Create the architecture Word document."""
    doc = Document()

    # Title
    title = doc.add_heading("EconStats Architecture", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph("Economic Data Query & Visualization System")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor(107, 114, 128)

    doc.add_paragraph()

    # Overview
    doc.add_heading("Overview", level=1)
    doc.add_paragraph(
        "EconStats is an AI-powered economic data visualization tool. Users ask questions "
        "in natural language (like \"How is inflation doing?\" or \"Compare US and Eurozone GDP\"), "
        "and the system automatically finds the right data series, fetches the data, and "
        "generates charts with AI-written summaries."
    )

    # Visual flow
    doc.add_heading("Query Flow Diagram", level=1)

    flow = doc.add_paragraph()
    flow.alignment = WD_ALIGN_PARAGRAPH.CENTER
    flow_text = """
┌─────────────────┐
│   User Query    │
│ "How is GDP?"   │
└────────┬────────┘
         ▼
┌─────────────────┐
│  Preprocessing  │
│ Extract dates,  │
│ demographics    │
└────────┬────────┘
         ▼
┌─────────────────┐
│  Query Router   │
│ US? Intl? Both? │
└────────┬────────┘
         ▼
┌─────────────────┐
│  Find Series    │
│ Plans/RAG/FRED  │
└────────┬────────┘
         ▼
┌─────────────────┐
│ LLM Validation  │
│ Is this right?  │
└────────┬────────┘
         ▼
┌─────────────────┐
│  Fetch Data     │
│ FRED + DBnomics │
└────────┬────────┘
         ▼
┌─────────────────┐
│   Display       │
│ Charts + AI     │
└─────────────────┘
"""
    run = flow.add_run(flow_text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)

    doc.add_page_break()

    # Detailed Steps
    doc.add_heading("Detailed Step-by-Step Explanation", level=1)

    # Step 1
    add_step_box(
        doc, 1, "User Query Input",
        "Accept a natural language question about economics from the user.",
        "The Streamlit web interface provides a text input where users type questions. "
        "The system handles questions ranging from simple (\"unemployment rate\") to complex "
        "(\"how has US growth compared to the Eurozone since COVID?\").",
        "\"How are Black workers doing in the labor market?\"",
        "app.py (Streamlit interface)"
    )

    # Step 2
    add_step_box(
        doc, 2, "Query Preprocessing",
        "Extract important context from the query before searching for data.",
        "Four extractors run on every query:\n"
        "• Temporal Extraction: Finds date references (\"in 2022\", \"pre-COVID\", \"last 5 years\")\n"
        "• Demographic Extraction: Identifies demographic groups (\"Black\", \"Hispanic\", \"women\")\n"
        "• Geographic Detection: Spots state/region mentions (\"Texas\", \"California\")\n"
        "• Synonym Mapping: Normalizes terms (\"jobs\" → \"employment\", \"pay\" → \"wages\")",
        "Query \"Black unemployment in Texas since COVID\" extracts:\n"
        "  → demographic: \"black\"\n"
        "  → geographic: \"texas\"\n"
        "  → temporal: 2020-03 to present",
        "app.py: extract_temporal_filter(), extract_demographic_group(), detect_geographic_scope()"
    )

    # Step 3
    add_step_box(
        doc, 3, "Smart Query Router",
        "Determine if the query needs US data, international data, or both.",
        "The router analyzes the query for:\n"
        "• Comparison keywords: \"vs\", \"compared to\", \"compare\"\n"
        "• Region mentions: US, Eurozone, UK, China, Japan, etc.\n"
        "• Indicator type: GDP, inflation, unemployment, interest rates\n\n"
        "Routes to: FRED only (US), DBnomics only (international), or both (comparisons).",
        "\"US vs Eurozone GDP\" → is_comparison=True → fetch from FRED + DBnomics",
        "agents/query_router.py"
    )

    # Step 4
    add_step_box(
        doc, 4, "Find Relevant Data Series",
        "Identify which specific data series (like UNRATE, GDPC1) answer the user's question.",
        "Three methods work together:\n\n"
        "1. Pre-computed Plans (350+ queries): Exact matches for common questions. "
        "\"how is inflation doing\" → [CPIAUCSL, CPILFESL, PCEPI]\n\n"
        "2. RAG Catalog (115+ series): Semantic search using TF-IDF to find series "
        "whose descriptions match the query.\n\n"
        "3. FRED API Search: Dynamic search for series not in our catalog, especially "
        "state-level data (TXUR for Texas unemployment).",
        "Query \"housing market\" matches plan → [HOUST, PERMIT, CSUSHPISA, MORTGAGE30US]",
        "agents/plans_*.json, agents/series_rag.py, FRED API"
    )

    doc.add_page_break()

    # Step 5
    add_step_box(
        doc, 5, "LLM Ensemble Validation",
        "Use AI to verify the selected series actually answer the user's question.",
        "An ensemble of 3 LLMs (Claude, Gemini, GPT) validates series:\n\n"
        "• Relevance Check: Does \"Manufacturing Employment\" answer \"solar energy jobs\"? → REJECT\n"
        "• Demographic Check: Does \"Women's Employment\" answer \"Black workers\"? → REJECT\n"
        "• Presentation Check: Is this a stock (show change), flow (show level), or rate?\n\n"
        "This prevents returning wrong data. If all series are rejected, the system "
        "shows a helpful \"no data available\" message instead of wrong data.",
        "Query \"Black workers\" with series \"Women's Employment\" → REJECTED (wrong demographic)",
        "agents/agent_ensemble.py: validate_series_relevance(), validate_presentation()"
    )

    # Step 6
    add_step_box(
        doc, 6, "Fetch Data from APIs",
        "Retrieve actual data values from external APIs.",
        "Three data sources:\n\n"
        "1. FRED API (US data): Federal Reserve Economic Data\n"
        "   - GDP, unemployment, inflation, payrolls, interest rates\n"
        "   - 800,000+ time series\n\n"
        "2. DBnomics API (International): Aggregates IMF, Eurostat, ECB, Bank of England\n"
        "   - Eurozone, UK, Japan, China, Germany GDP/inflation\n"
        "   - Central bank rates\n\n"
        "3. Polymarket API (Forward-looking): Prediction market data\n"
        "   - Recession probability, Fed rate expectations",
        "Comparison query fetches GDPC1 from FRED + eurozone_gdp from DBnomics",
        "FRED API, agents/dbnomics.py, agents/polymarket.py"
    )

    # Step 7
    add_step_box(
        doc, 7, "Transform Data for Display",
        "Convert raw data into the right format for meaningful display.",
        "Different series need different treatments:\n\n"
        "• STOCK (cumulative total like GDP in $billions): Transform to YoY % change\n"
        "  Raw: $28.3 trillion → Display: +2.3% YoY\n\n"
        "• FLOW (per-period like weekly claims): Show as level\n"
        "  Raw: 220,000 → Display: 220K\n\n"
        "• RATE (already a percentage): Show as level\n"
        "  Raw: 4.1% → Display: 4.1%",
        "GDPC1 (Real GDP) level $28T → transformed to \"GDP grew 2.3% year-over-year\"",
        "app.py: calculate_yoy_growth(), validate_presentation()"
    )

    # Step 8
    add_step_box(
        doc, 8, "Generate Visualization & Summary",
        "Create charts and AI-written analysis for the user.",
        "The presentation layer includes:\n\n"
        "• Interactive Charts (Altair): Line charts with proper axis labels, "
        "recession shading, multi-series support\n\n"
        "• AI Summary: LLM-generated 2-3 sentence summary of what the data shows, "
        "highlighting trends and notable changes\n\n"
        "• Polymarket Predictions: For relevant queries, shows forward-looking "
        "predictions (e.g., \"Recession in 2025: 23% probability\")",
        "GDP query shows chart + \"US GDP grew 2.3% YoY in Q3 2024, driven by consumer spending...\"",
        "app.py (Streamlit + Altair)"
    )

    doc.add_page_break()

    # Critical Rules
    doc.add_heading("Critical Validation Rules", level=1)

    p = doc.add_paragraph()
    run = p.add_run("The system enforces strict rules to prevent showing wrong data:")
    run.bold = True

    # Comparison rules table
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'

    headers = ["Rule", "What it Prevents", "Example"]
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True

    rules = [
        ("YoY vs YoY only", "Comparing 2.3% annual growth to 0.6% quarterly growth",
         "US GDP YoY vs Eurozone GDP QoQ = BLOCKED"),
        ("Real vs Real only", "Comparing inflation-adjusted to nominal values",
         "Real GDP vs Nominal GDP = BLOCKED"),
        ("Demographic matching", "Returning women's data for query about Black workers",
         "\"Black unemployment\" → Women's employment = BLOCKED"),
    ]

    for i, (rule, prevents, example) in enumerate(rules):
        row = table.rows[i + 1]
        row.cells[0].text = rule
        row.cells[1].text = prevents
        row.cells[2].text = example

    doc.add_paragraph()

    # File Reference
    doc.add_heading("Key Files Reference", level=1)

    files_table = doc.add_table(rows=11, cols=2)
    files_table.style = 'Table Grid'

    files_table.rows[0].cells[0].text = "File"
    files_table.rows[0].cells[1].text = "Purpose"
    files_table.rows[0].cells[0].paragraphs[0].runs[0].bold = True
    files_table.rows[0].cells[1].paragraphs[0].runs[0].bold = True

    files = [
        ("app.py", "Main application - Streamlit UI, query orchestration, chart generation"),
        ("agents/agent_ensemble.py", "LLM ensemble (Claude + Gemini + GPT) for validation"),
        ("agents/series_rag.py", "RAG catalog with 115+ curated FRED series"),
        ("agents/query_router.py", "Smart routing for US vs international vs comparison queries"),
        ("agents/dbnomics.py", "International data from IMF, Eurostat, ECB, Bank of England"),
        ("agents/polymarket.py", "Prediction market data (recession odds, Fed expectations)"),
        ("agents/stocks.py", "Stock market query plans (S&P 500, VIX, etc.)"),
        ("agents/plans_*.json", "350+ pre-computed query-to-series mappings"),
        ("CLAUDE.local.md", "Project memory - rules and patterns for AI assistants"),
        ("requirements.txt", "Python dependencies"),
    ]

    for i, (file, purpose) in enumerate(files):
        row = files_table.rows[i + 1]
        run = row.cells[0].paragraphs[0].add_run(file)
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        row.cells[1].text = purpose

    doc.add_paragraph()

    # Data Sources
    doc.add_heading("Data Sources", level=1)

    sources_table = doc.add_table(rows=4, cols=3)
    sources_table.style = 'Table Grid'

    sources_table.rows[0].cells[0].text = "Source"
    sources_table.rows[0].cells[1].text = "Coverage"
    sources_table.rows[0].cells[2].text = "Examples"
    for cell in sources_table.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True

    sources = [
        ("FRED API", "US economic data (800K+ series)", "GDP, unemployment, inflation, payrolls"),
        ("DBnomics API", "International (IMF, Eurostat, ECB, BOE)", "Eurozone GDP, UK inflation, ECB rate"),
        ("Polymarket API", "Prediction markets", "Recession odds, Fed rate expectations"),
    ]

    for i, (source, coverage, examples) in enumerate(sources):
        row = sources_table.rows[i + 1]
        row.cells[0].text = source
        row.cells[1].text = coverage
        row.cells[2].text = examples

    # Save
    doc.save(filename)
    print(f"Word document saved to: {filename}")

if __name__ == "__main__":
    output_path = "/Users/josh/Desktop/econstats/EconStats_Architecture.docx"
    create_architecture_doc(output_path)
